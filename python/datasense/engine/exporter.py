import os
import json
from datetime import datetime
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class ReportExporter:
    """Exports dataset analysis into professional reports with watermarking."""

    WATERMARK = "Made with DataSense: https://github.com/jeffasante/datasense"

    def export_markdown(self, data: Dict[str, Any], output_path: str):
        """Generates a comprehensive Markdown report."""
        fingerprints = data.get("fingerprints", {})
        leaderboard = data.get("leaderboard", [])
        plots = data.get("plots", {})
        recommendation = data.get("recommendation", {})
        summary = data.get("summary", "")

        md = []
        md.append("# DataSense Analysis Report")
        md.append(f"*Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*")
        md.append("\n---\n")

        # 1. Executive Summary
        md.append("## Executive Summary")
        if summary:
            md.append(summary)
            md.append("\n")
        
        if recommendation:
            primary = recommendation.get("primary", "N/A")
            conf = recommendation.get("confidence") or recommendation.get("score", 0)
            score_formatted = int(conf * 100) if isinstance(conf, (int, float)) else 0
            md.append(f"**Primary Architecture**: `{primary}` ({score_formatted}% match)")
        md.append("\n")

        # 2. Visualizations
        if plots:
            md.append("## Visual Insights")
            for key, plot_path in plots.items():
                if plot_path:
                    base_dir = os.path.dirname(output_path)
                    try: rel_plot = os.path.relpath(plot_path, base_dir)
                    except: rel_plot = plot_path
                    md.append(f"### {key.replace('_', ' ').capitalize()}\n![]({rel_plot})\n")

        # 3. Leaderboard
        if leaderboard:
            md.append("## Leaderboard Ranking")
            md.append("| Rank | Model | Score | Justification |")
            md.append("| :--- | :--- | :--- | :--- |")
            for i, entry in enumerate(leaderboard[:10]):
                score = f"{int(entry.get('score', 0) * 100)}%"
                md.append(f"| {i+1} | **{entry.get('model', '?')}** | {score} | {entry.get('justification', '')} |")
            md.append("\n")

        # Watermark
        md.append("\n---\n")
        md.append(f"*{self.WATERMARK}*")

        with open(output_path, "w") as f:
            f.write("\n".join(md))
        return output_path

    def export_txt(self, data: Dict[str, Any], output_path: str):
        """Generates a plain text report with watermark."""
        fingerprints = data.get("fingerprints", {})
        leaderboard = data.get("leaderboard", [])
        recommendation = data.get("recommendation", {})
        
        txt = []
        txt.append("DATASENSE ANALYSIS REPORT")
        txt.append("=" * 30)
        txt.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        
        if recommendation:
            txt.append(f"PRIMARY MODEL: {recommendation.get('primary', 'N/A')}")
            conf = int((recommendation.get("confidence") or recommendation.get("score", 0)) * 100)
            txt.append(f"MATCH SCORE: {conf}%\n")
        
        if leaderboard:
            txt.append("LEADERBOARD")
            txt.append("-" * 30)
            for i, entry in enumerate(leaderboard[:10]):
                s = int(entry.get('score', 0) * 100)
                txt.append(f"{i+1}. {entry.get('model', '?')} ({s}%)")
            txt.append("\n")
            
        txt.append("-" * 30)
        txt.append(self.WATERMARK)

        with open(output_path, "w") as f:
            f.write("\n".join(txt))
        return output_path

    def export_docx(self, data: Dict[str, Any], output_path: str):
        """Generates a professional DOCX report with footer watermark."""
        if Document is None:
            return self.export_markdown(data, output_path.replace(".docx", ".md"))

        doc = Document()
        doc.add_heading("DataSense Analysis Report", 0)
        
        # Main content
        doc.add_heading("Executive Summary", level=1)
        rec = data.get("recommendation", {})
        if rec:
            p = doc.add_paragraph()
            p.add_run("Primary Recommendation: ").bold = True
            p.add_run(rec.get("primary", "N/A"))

        plots = data.get("plots", {})
        if plots:
            doc.add_heading("Visual Insights", level=1)
            for key, p_path in plots.items():
                if os.path.exists(p_path):
                    doc.add_heading(key.replace('_', ' ').capitalize(), level=2)
                    doc.add_picture(p_path, width=Inches(5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        leaderboard = data.get("leaderboard", [])
        if leaderboard:
            doc.add_heading("Architecture Ranking", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Model'
            hdr[1].text = 'Score'
            hdr[2].text = 'Justification'
            for entry in leaderboard[:10]:
                row = table.add_row().cells
                row[0].text = entry.get("model", "?")
                row[1].text = f"{int(entry.get('score', 0) * 100)}%"
                row[2].text = entry.get("justification", "")

        # Watermark Footer
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = self.WATERMARK
        p.style.font.size = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_path)
        return output_path

    def export_pdf(self, data: Dict[str, Any], output_path: str):
        """Generates a professional PDF report with footer watermark."""
        if not PDF_AVAILABLE:
            return self.export_txt(data, output_path.replace(".pdf", ".txt"))

        class PDF(FPDF):
            def footer(self):
                self.set_y(-15)
                self.set_font("helvetica", "I", 8)
                self.set_text_color(100, 100, 100)
                self.cell(0, 10, ReportExporter.WATERMARK, 0, 0, "C")

        pdf = PDF()
        pdf.add_page()
        pdf.set_font("helvetica", "B", 16)
        pdf.cell(0, 10, "DataSense Analysis Report", ln=True, align="C")
        pdf.set_font("helvetica", "I", 10)
        pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", ln=True, align="C")
        pdf.ln(10)

        # Summary
        rec = data.get("recommendation", {})
        if rec:
            pdf.set_font("helvetica", "B", 14)
            pdf.cell(0, 10, "Strategic Recommendation", ln=True)
            pdf.set_font("helvetica", "", 11)
            primary = rec.get("primary", "N/A")
            pdf.multi_cell(0, 8, f"Based on the unique fingerprint of your dataset, the optimal architecture for deployment is {primary}.\n")
            pdf.ln(5)

        # Plots
        plots = data.get("plots", {})
        if plots:
            pdf.add_page()  # Start visuals on new page
            pdf.set_font("helvetica", "B", 14)
            pdf.cell(0, 10, "Deep Visual Insights", ln=True)
            pdf.ln(5)
            
            # Prioritize leaderboard and fingerprint
            for key in ["leaderboard", "fingerprint", "audio_profile", "image_profile", "tabular_profile"]:
                if key in plots and os.path.exists(plots[key]):
                    pdf.set_font("helvetica", "B", 11)
                    pdf.cell(0, 10, key.replace('_', ' ').capitalize(), ln=True)
                    pdf.image(plots[key], w=170)
                    pdf.ln(10)

        pdf.output(output_path)
        return output_path
